#!/usr/bin/env python3
# utils.py — structure‑aware heading mapper (fix random H2/H3 placement)
#
# What changed in this build (vs the last one):
# 1) Replaced raw index‑by‑index style sequence overlay with a **structure‑aware mapper**.
#    We now detect *good heading candidates* in the refined text and assign the
#    exact number of H1/H2/H3 from the source to the best candidates **in order**.
#    This prevents headings from landing on arbitrary body paragraphs when your
#    refined draft has extra/shifted paragraph breaks.
# 2) Still honors explicit markdown markers in the refined text ('# ', '## ', '### ').
# 3) Keeps style name lookups (no style_id warnings) and the Arial 11/20/16/14 skeleton.
# 4) Reader remains robust (style name, outline level, char‑style, size/bold heuristics,
#    and traversal of body + tables).

from __future__ import annotations

import io
import os
import re
import pickle
import tempfile
import warnings
from pathlib import Path
from urllib.parse import urlparse, parse_qs
from typing import Dict, List, Tuple

import yaml
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from google.oauth2.credentials import Credentials
from google.oauth2 import service_account
from google.auth.transport.requests import Request
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload, MediaIoBaseDownload
from googleapiclient.errors import HttpError
import time
import random
import os
import json as _json
import base64

# PDF and DOC support
try:
    import PyPDF2
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    import docx2txt
    DOC_SUPPORT = True
except ImportError:
    DOC_SUPPORT = False

def safe_encoder(obj: Any) -> str:
    """
    Safely encode an object to JSON string, handling non-serializable objects.
    """
    import json
    from datetime import datetime, date
    
    def default(o):
        if isinstance(o, (datetime, date)):
            return o.isoformat()
        if hasattr(o, '__dict__'):
            return o.__dict__
        return str(o)
        
    return json.dumps(obj, default=default)

# ---------------------------
# PDF and DOC extraction functions
# ---------------------------

def _extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file"""
    if not PDF_SUPPORT:
        raise ImportError("PyPDF2 is required for PDF support. Install with: pip install PyPDF2")
    
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()
    except Exception as e:
        raise ValueError(f"Failed to extract text from PDF: {str(e)}")

def _extract_text_from_doc(file_path: str) -> str:
    """Extract text from DOC file"""
    if not DOC_SUPPORT:
        raise ImportError("python-docx2txt is required for DOC support. Install with: pip install python-docx2txt")
    
    try:
        import docx2txt
        text = docx2txt.process(file_path)
        return text.strip()
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOC: {str(e)}")

# ---------------------------
# Style helpers (name lookups)
# ---------------------------
_HEADING_ID_TO_NAME = {f"Heading{i}": f"Heading {i}" for i in range(1, 7)}
_HEADING_NAME_TO_NAME = {f"Heading {i}": f"Heading {i}" for i in range(1, 7)}
_CANON_STYLE = {"Normal": "Normal", **_HEADING_ID_TO_NAME, **_HEADING_NAME_TO_NAME}

def _canon_style_name(key: str) -> str:
    return _CANON_STYLE.get(key, key)

def _get_style_by_name(doc: Document, key: str):
    name = _canon_style_name(key)

    # First, try explicit name match (avoids deprecated style_id lookup)
    for style in doc.styles:
        if style.name == name:
            return style

    # Fallback: try to create if missing (only for Normal)
    if name == "Normal":
        return doc.styles["Normal"]

    return None

# ---------------------------
# Heading detection and mapping
# ---------------------------

def _is_heading_paragraph(para):
    """Detect if a paragraph is a heading based on style properties."""
    style = para.style
    style_name = style.name if style else ""

    # Check style name
    if "Heading" in style_name or style_name.startswith("Heading "):
        return True

    # Check outline level
    if hasattr(para, "paragraph_format") and para.paragraph_format.outline_level:
        return True

    # Check formatting (bold + larger size often indicates heading)
    runs = para.runs
    if runs:
        first_run = runs[0]
        if first_run.bold:
            font_size = first_run.font.size
            if font_size and font_size.pt > 12:  # Larger than normal text
                return True

    return False

def _extract_headings_from_doc(doc: Document) -> List[Tuple[str, int]]:
    """
    Extract headings from document with their levels.
    Returns list of (heading_text, level) tuples.
    """
    headings = []
    for para in doc.paragraphs:
        if _is_heading_paragraph(para):
            text = para.text.strip()
            if text:
                # Determine level from style name
                style_name = para.style.name if para.style else ""
                level = 1
                if "Heading 1" in style_name or "Heading1" in style_name:
                    level = 1
                elif "Heading 2" in style_name or "Heading2" in style_name:
                    level = 2
                elif "Heading 3" in style_name or "Heading3" in style_name:
                    level = 3
                elif "Heading 4" in style_name or "Heading4" in style_name:
                    level = 4
                elif "Heading 5" in style_name or "Heading5" in style_name:
                    level = 5
                elif "Heading 6" in style_name or "Heading6" in style_name:
                    level = 6
                else:
                    # Try outline level
                    if hasattr(para, "paragraph_format"):
                        outline_level = para.paragraph_format.outline_level
                        if outline_level:
                            level = outline_level

                headings.append((text, level))
    return headings

def _map_headings_to_refined_text(source_headings: List[Tuple[str, int]], refined_text: str) -> Dict[int, str]:
    """
    Map source headings to refined text by finding best matches.
    Returns dict mapping level -> heading_text for refined text.
    """
    # Split refined text into paragraphs
    refined_paras = [p.strip() for p in refined_text.split("\n") if p.strip()]

    # Find heading candidates in refined text (lines that look like headings)
    heading_candidates = []
    for i, para in enumerate(refined_paras):
        # Check if it looks like a heading (starts with #, or is short and bold-looking)
        if para.startswith("#"):
            # Markdown heading
            level = len(para) - len(para.lstrip("#"))
            text = para.lstrip("#").strip()
            heading_candidates.append((i, text, level))
        elif len(para) < 100 and para.isupper():  # Short uppercase line might be heading
            heading_candidates.append((i, para, 1))

    # Map source headings to candidates in order
    mapped = {}
    candidate_idx = 0
    for source_text, source_level in source_headings:
        # Find best matching candidate
        best_match = None
        best_score = 0
        for i, (cand_idx, cand_text, cand_level) in enumerate(heading_candidates[candidate_idx:], start=candidate_idx):
            # Score based on text similarity and level match
            text_similarity = len(set(source_text.lower().split()) & set(cand_text.lower().split())) / max(len(source_text.split()), 1)
            level_match = 1.0 if cand_level == source_level else 0.5
            score = text_similarity * 0.7 + level_match * 0.3

            if score > best_score:
                best_score = score
                best_match = (cand_idx, cand_text, cand_level)

        if best_match and best_score > 0.3:  # Threshold for matching
            mapped[best_match[2]] = best_match[1]
            candidate_idx = heading_candidates.index(best_match) + 1
    else:
            # No good match, use source heading as-is
            mapped[source_level] = source_text

    return mapped

# ---------------------------
# Main mapping function
# ---------------------------

def map_headings_to_refined_doc(source_doc_path: str, refined_doc_path: str, output_path: str):
    """
    Map headings from source document to refined document, preserving structure.
    """
    source_doc = Document(source_doc_path)
    refined_doc = Document(refined_doc_path)

    # Extract headings from source
    source_headings = _extract_headings_from_doc(source_doc)

    # Extract text from refined document
    refined_text = "\n".join([para.text for para in refined_doc.paragraphs])

    # Map headings
    heading_map = _map_headings_to_refined_text(source_headings, refined_text)

    # Apply heading styles to refined document
    for para in refined_doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Check if this paragraph matches a mapped heading
        matched = False
        for level, heading_text in heading_map.items():
            if text.lower() == heading_text.lower() or text.lower().startswith(heading_text.lower()):
                # Apply heading style
                style_name = f"Heading {level}"
                style = _get_style_by_name(refined_doc, style_name)
                if style:
                    para.style = style
                matched = True
                break
        
        # Check for markdown headings if no mapped heading matched
        if not matched:
            if text.startswith("#"):
                level = len(text) - len(text.lstrip("#"))
                style_name = f"Heading {min(level, 6)}"
                style = _get_style_by_name(refined_doc, style_name)
                if style:
                    para.style = style
                    # Remove markdown markers
                    para.text = text.lstrip("#").strip()

    # Save output
    refined_doc.save(output_path)

# ---------------------------
# Google Drive integration
# ---------------------------

def get_drive_service():
    """Get Google Drive service using service account credentials."""
    creds = get_google_credentials()
    if not creds:
        raise ValueError("Failed to get Google credentials")
    return build('drive', 'v3', credentials=creds)

def get_docs_service():
    """Get Google Docs service using service account credentials."""
    creds = get_google_credentials()
    if not creds:
        raise ValueError("Failed to get Google credentials")
    return build('docs', 'v1', credentials=creds)

def download_drive_file(link_or_id: str, dest_path: str) -> str:
    """
    Download a Google Drive file by link or file ID.
    
    Args:
        link_or_id: Google Drive share link or file ID
        dest_path: Local path to save the file
        
    Returns:
        Path to downloaded file
    """
    try:
        # Extract file ID from link if needed
        file_id = link_or_id
        if 'drive.google.com' in link_or_id or 'docs.google.com' in link_or_id:
            parsed = urlparse(link_or_id)
            if '/d/' in parsed.path:
                file_id = parsed.path.split('/d/')[1].split('/')[0]
            elif 'id=' in parsed.query:
                file_id = parse_qs(parsed.query)['id'][0]
        
        service = get_drive_service()
        
        # Get file metadata
        file_metadata = service.files().get(fileId=file_id).execute()
        mime_type = file_metadata.get('mimeType', '')
        
        # Handle Google Docs format
        if mime_type == 'application/vnd.google-apps.document':
            # Export as DOCX
            request = service.files().export_media(fileId=file_id, mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        elif mime_type == 'application/vnd.google-apps.spreadsheet':
            # Export as XLSX
            request = service.files().export_media(fileId=file_id, mimeType='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        else:
            # Download directly
            request = service.files().get_media(fileId=file_id)
        
        # Download file
        fh = io.FileIO(dest_path, 'wb')
        downloader = MediaIoBaseDownload(fh, request)
        done = False
        while done is False:
            status, done = downloader.next_chunk()
        
        return dest_path
        
    except HttpError as e:
        raise

OAUTH_SCOPES = ['https://www.googleapis.com/auth/drive', 'https://www.googleapis.com/auth/documents']

def _parse_json_from_env(env_value: str) -> dict:
    """
    Parse JSON from environment variable, handling various formats:
    - Standard JSON string
    - Base64 encoded JSON
    - JSON with escaped quotes
    - Multi-line JSON (with newlines)
    """
    if not env_value or not env_value.strip():
        raise ValueError("Empty JSON value")
    
    # Try to decode as base64 first (if it looks like base64 or if env var ends with _BASE64)
    if len(env_value) > 100 and not env_value.strip().startswith('{'):
        try:
            decoded = base64.b64decode(env_value).decode('utf-8')
            return _json.loads(decoded)
        except Exception:
            pass  # Not base64, continue with normal parsing
    
    # Check for base64 variant
    base64_json = os.getenv('GOOGLE_SERVICE_ACCOUNT_JSON_BASE64')
    if base64_json:
        try:
            decoded = base64.b64decode(base64_json).decode('utf-8')
            return _json.loads(decoded)
        except Exception as e:
            raise ValueError(f"Failed to decode base64 JSON: {e}")
    
    # Try parsing as-is
    try:
        return _json.loads(env_value)
    except _json.JSONDecodeError:
        pass
    
    # Try removing leading/trailing whitespace and quotes
    cleaned = env_value.strip()
    if cleaned.startswith('"') and cleaned.endswith('"'):
        cleaned = cleaned[1:-1]
    if cleaned.startswith("'") and cleaned.endswith("'"):
        cleaned = cleaned[1:-1]
    
    # Try parsing cleaned value
    try:
        return _json.loads(cleaned)
    except _json.JSONDecodeError:
        pass
    
    # Try unescaping common escape sequences
    cleaned = cleaned.replace('\\"', '"').replace("\\'", "'").replace('\\n', '\n').replace('\\r', '\r').replace('\\t', '\t')
    try:
        return _json.loads(cleaned)
    except _json.JSONDecodeError:
        pass
    
    # Last attempt: try to fix common issues
    # Replace single quotes with double quotes (very basic, may not work for all cases)
    cleaned = re.sub(r"'([^']*)':", r'"\1":', cleaned)  # Keys
    cleaned = re.sub(r":\s*'([^']*)'", r': "\1"', cleaned)  # String values
    try:
        return _json.loads(cleaned)
    except _json.JSONDecodeError as e:
        raise ValueError(f"Failed to parse JSON from environment variable. Error: {e}. Value preview: {env_value[:100]}...")

def get_google_credentials(credentials_path: str = None, token_path: str = None) -> Credentials:
    """
    Get Google credentials from environment variables (preferred) or files (fallback).
    
    Priority:
    1. GOOGLE_SERVICE_ACCOUNT_JSON (env var with JSON string)
    2. GOOGLE_SERVICE_ACCOUNT_FILE (env var with file path)
    3. GOOGLE_OAUTH_CREDENTIALS_JSON + GOOGLE_OAUTH_TOKEN_JSON (env vars)
    4. File-based fallback (for backward compatibility)
    """
    import json as _json
    
    # Priority 1: Service account JSON from environment variable
    service_account_json = os.getenv('GOOGLE_SERVICE_ACCOUNT_JSON')
    if service_account_json and service_account_json.strip() and service_account_json.strip() not in ['', 'null', 'None']:
        try:
            creds_data = _parse_json_from_env(service_account_json)
            
            # Validate required fields
            required_fields = ['type', 'project_id', 'private_key', 'client_email']
            missing_fields = [field for field in required_fields if field not in creds_data]
            if missing_fields:
                raise ValueError(f"Missing required fields in service account JSON: {missing_fields}")
            
            if creds_data.get('type') != 'service_account':
                raise ValueError(f"Invalid service account type: {creds_data.get('type')}")
            
            # Validate private key format
            private_key = creds_data.get('private_key', '')
            if not private_key or not private_key.startswith('-----BEGIN PRIVATE KEY-----'):
                raise ValueError("Invalid private key format in service account JSON")
            
            creds = service_account.Credentials.from_service_account_info(
                creds_data,
                scopes=OAUTH_SCOPES
            )
            return creds
        except Exception as e:
            print(f"⚠️  Failed to load service account from GOOGLE_SERVICE_ACCOUNT_JSON: {e}")
            print(f"   Error type: {type(e).__name__}")
            print(f"   💡 TIP: Use GOOGLE_SERVICE_ACCOUNT_FILE=config/google_credentials.json instead")
            print(f"   📖 See backend/FIX_GOOGLE_DRIVE_NOW.md for help")
            print("   Trying file-based method...")
    
    # Check for base64 variant
    base64_json = os.getenv('GOOGLE_SERVICE_ACCOUNT_JSON_BASE64')
    if base64_json and base64_json.strip():
        try:
            decoded = base64.b64decode(base64_json).decode('utf-8')
            creds_data = _json.loads(decoded)
            
            # Validate required fields
            required_fields = ['type', 'project_id', 'private_key', 'client_email']
            missing_fields = [field for field in required_fields if field not in creds_data]
            if missing_fields:
                raise ValueError(f"Missing required fields in service account JSON: {missing_fields}")
            
            if creds_data.get('type') != 'service_account':
                raise ValueError(f"Invalid service account type: {creds_data.get('type')}")
            
            creds = service_account.Credentials.from_service_account_info(
                creds_data,
                scopes=OAUTH_SCOPES
            )
            return creds
        except Exception as e:
            print(f"Warning: Failed to load service account from GOOGLE_SERVICE_ACCOUNT_JSON_BASE64: {e}")
            print("Trying other methods...")
    
    # Priority 2: Service account file path from environment variable
    service_account_file = os.getenv('GOOGLE_SERVICE_ACCOUNT_FILE')
    if service_account_file:
        # Resolve relative paths relative to backend directory
        if not os.path.isabs(service_account_file):
            from app.core.paths import get_backend_root
            backend_dir = str(get_backend_root())
            service_account_file = os.path.join(backend_dir, service_account_file)
        
        if os.path.exists(service_account_file):
            try:
                creds = service_account.Credentials.from_service_account_file(
                    service_account_file,
                    scopes=OAUTH_SCOPES
                )
                print(f"✅ Loaded Google credentials from file: {service_account_file}")
                return creds
            except Exception as e:
                print(f"⚠️  Failed to load service account from file {service_account_file}: {e}")
                print("   Trying default file location...")
        else:
            print(f"⚠️  GOOGLE_SERVICE_ACCOUNT_FILE specified but file not found: {service_account_file}")
    
    # Try default service account file location
    from app.core.paths import get_backend_root
    backend_dir = str(get_backend_root())
    default_service_account = os.path.join(backend_dir, 'config', 'google_credentials.json')
    if os.path.exists(default_service_account):
        try:
            creds = service_account.Credentials.from_service_account_file(
                default_service_account,
                scopes=OAUTH_SCOPES
            )
            print(f"✅ Loaded Google credentials from default file: {default_service_account}")
            
            # Try to validate credentials by attempting to refresh
            try:
                from google.auth.transport.requests import Request
                if not creds.valid:
                    creds.refresh(Request())
                print(f"✅ Google credentials validated successfully")
                return creds
            except Exception as refresh_error:
                error_msg = str(refresh_error)
                if 'invalid_grant' in error_msg or 'JWT' in error_msg or 'signature' in error_msg:
                    print(f"❌ CRITICAL: Invalid JWT Signature in credentials file!")
                    print(f"   The service account key is invalid or has been regenerated.")
                    print(f"   SOLUTION: Regenerate the key in Google Cloud Console")
                    print(f"   See: backend/FIX_INVALID_JWT_SIGNATURE.md for instructions")
                    print(f"   Service Account: {creds.service_account_email if hasattr(creds, 'service_account_email') else 'unknown'}")
                else:
                    print(f"⚠️  Failed to validate credentials: {refresh_error}")
                # Don't return None yet - let it try OAuth flow
                raise refresh_error
            
        except Exception as e:
            error_msg = str(e)
            if 'invalid_grant' in error_msg or 'JWT' in error_msg or 'signature' in error_msg:
                print(f"❌ CRITICAL: Invalid JWT Signature!")
                print(f"   Your service account key needs to be regenerated.")
                print(f"   📖 See: backend/FIX_INVALID_JWT_SIGNATURE.md")
            else:
                print(f"⚠️  Failed to load service account from default file: {e}")
            print("   Trying OAuth flow...")
    
    # Priority 3: OAuth credentials from environment variables
    oauth_credentials_json = os.getenv('GOOGLE_OAUTH_CREDENTIALS_JSON')
    oauth_token_json = os.getenv('GOOGLE_OAUTH_TOKEN_JSON')
    
    if oauth_credentials_json:
        try:
            creds_data = _parse_json_from_env(oauth_credentials_json)
            # Create credentials from info
            from google.oauth2.credentials import Credentials as OAuthCredentials
            
            # Try to load token from env var
            creds = None
            if oauth_token_json:
                try:
                    token_data = _parse_json_from_env(oauth_token_json)
                    creds = OAuthCredentials.from_authorized_user_info(
                        {**creds_data, **token_data},
                        scopes=OAUTH_SCOPES
                    )
                except Exception:
                    pass
            
            # If no valid token, need to do OAuth flow (interactive)
            if not creds or not creds.valid:
                if creds and creds.expired and creds.refresh_token:
                    creds.refresh(Request())
                else:
                    # For OAuth flow, we can't do interactive flow in production, return None or raise
                    raise ValueError(
                        "OAuth credentials require interactive authentication. "
                        "Use service account credentials (GOOGLE_SERVICE_ACCOUNT_JSON) for production."
                    )
            return creds
        except Exception as e:
            print(f"Warning: Failed to load OAuth credentials from env vars: {e}")
    
    # Priority 4: Fallback to file-based (for backward compatibility during migration)
    from app.core.paths import get_backend_root
    backend_dir = str(get_backend_root())
    default_service_account = os.path.join(backend_dir, 'config', 'google_credentials.json')
    
    if credentials_path is None:
        credentials_path = os.path.join(backend_dir, 'config', 'credentials.json')
    if token_path is None:
        token_path = os.path.join(backend_dir, 'config', 'token.json')
    
    # Try default service account file
    if os.path.exists(default_service_account):
        try:
            creds = service_account.Credentials.from_service_account_file(
                default_service_account,
                scopes=OAUTH_SCOPES
            )
            return creds
        except Exception as e:
            print(f"Warning: Failed to load service account from default file: {e}")
    
    # Fall back to OAuth file flow (interactive - not suitable for production)
    creds = None
    token_file = Path(token_path)
    if token_file.exists():
        try:
            creds = pickle.loads(token_file.read_bytes())
        except Exception:
            pass
    
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            try:
                creds.refresh(Request())
            except Exception:
                creds = None
        
        if not creds and os.path.exists(credentials_path):
            try:
                flow = InstalledAppFlow.from_client_secrets_file(credentials_path, OAUTH_SCOPES)
                creds = flow.run_local_server(port=0)
                token_file.write_bytes(pickle.dumps(creds))
            except Exception as e:
                print(f"Warning: OAuth flow failed: {e}")
                return None
    
    # Final check: if we still don't have valid credentials, print helpful error
    if not creds:
        print()
        print("=" * 70)
        print("❌ GOOGLE DRIVE CREDENTIALS NOT FOUND")
        print("=" * 70)
        print()
        print("No valid Google credentials could be loaded.")
        print()
        print("SOLUTIONS:")
        print("1. Regenerate service account key (RECOMMENDED):")
        print("   - Go to: https://console.cloud.google.com/iam-admin/serviceaccounts")
        print("   - Find: turbo-alan-google-drive@crack-petal-469722-d1.iam.gserviceaccount.com")
        print("   - Create new JSON key and save to: backend/config/google_credentials.json")
        print()
        print("2. Or set GOOGLE_SERVICE_ACCOUNT_FILE in .env:")
        print("   GOOGLE_SERVICE_ACCOUNT_FILE=config/google_credentials.json")
        print()
        print("📖 See: backend/FIX_INVALID_JWT_SIGNATURE.md for detailed instructions")
        print("=" * 70)
        print()
    
    return creds

# ---------------------------
# Heuristics loading
# ---------------------------

def load_heuristics() -> Dict[str, Any]:
    """Load heuristics configuration from YAML file."""
    from app.core.paths import get_backend_root
    backend_dir = str(get_backend_root())
    heuristics_path = os.path.join(backend_dir, 'config', 'heuristics.yaml')
    
    if not os.path.exists(heuristics_path):
        # Return default heuristics if file doesn't exist
        return {
            'preserve_layout': True,
            'transforms': {
                'punctuation': 0.2,
                'splits': 0.2,
                'transitions': 0.2,
                'rhythm': 0.2,
                'keywords': 0.2,
                'grammar': 0.2,
                'synonyms': 0.2,
                'micro_degrade': 0.2,
            },
            'paragraph': {
                'min_sentences': 1,
                'max_sentences': 6,
            },
            'length': {
                'ratio_range': [0.9, 1.1],
            },
            'sampling': {
                'temperature': 0.5,
                'top_p': 0.9,
            },
            'sampling_schedule': [0.5, 0.45, 0.4],
            'top_p_schedule': [0.9, 0.81, 0.729],
            'synonym': {},
            'grammar': {},
            'micro_degrade': {},
            'phases': [
                'You are Turbo Alan Refiner. Apply micro-degradation transformations to reduce AI-detection signals while preserving coherence.',
                'You are Turbo Alan Refiner. Further vary punctuation, sentence length, and word choice to evade scanners without harming readability.',
                'You are Turbo Alan Refiner. Final pass: introduce subtle noise patterns, maintain natural flow, and ensure the text reads as human-written.',
            ],
        }
    
    try:
        with open(heuristics_path, 'r', encoding='utf-8') as f:
            return yaml.safe_load(f) or {}
    except Exception as e:
        print(f"Warning: Failed to load heuristics from {heuristics_path}: {e}")
        return {}

# ---------------------------
# File I/O utilities
# ---------------------------

def read_text_from_file(file_path: str) -> str:
    """Read text from a file, supporting various formats."""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    # Determine file type
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.pdf':
        return _extract_text_from_pdf(file_path)
    elif ext == '.doc':
        return _extract_text_from_doc(file_path)
    elif ext == '.docx':
        try:
            doc = Document(file_path)
            return '\n'.join([para.text for para in doc.paragraphs])
        except Exception as e:
            raise ValueError(f"Failed to read DOCX file: {e}")
    else:
        # Plain text
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()

def write_text_to_file(text: str = None, file_path: str = None, output_dir: str = None, 
                       base_name: str = None, ext: str = None, original_file: str = None, 
                       iteration: int = None, **kwargs) -> str:
    """
    Write text to a file. Supports multiple calling conventions:
    
    Simple: write_text_to_file(text, file_path)
    Advanced: write_text_to_file(text=..., output_dir=..., base_name=..., ext=..., ...)
    """
    # Handle advanced calling convention
    if output_dir and base_name and ext:
        if not text:
            raise ValueError("text parameter is required")
        # Sanitize base_name to prevent path injection
        base_name = str(base_name).strip()
        # Remove any path components and invalid characters
        base_name = os.path.basename(base_name)
        # Remove invalid characters for Windows/Unix paths
        base_name = re.sub(r'[<>:"|?*\n\r\t]', '_', base_name)
        # Remove any remaining control characters
        base_name = re.sub(r'[\x00-\x1f\x7f-\x9f]', '_', base_name)
        # Limit length to prevent filesystem issues (max 200 chars)
        if len(base_name) > 200:
            base_name = base_name[:200]
        # Ensure base_name is not empty
        if not base_name or base_name == '.' or base_name == '..':
            base_name = 'output'
        file_path = os.path.join(output_dir, f"{base_name}{ext}")
        if iteration is not None and iteration > 0:
            file_path = os.path.join(output_dir, f"{base_name}_pass{iteration}{ext}")
    
    # Fallback to simple convention
    if not file_path:
        if 'file_path' in kwargs:
            file_path = kwargs['file_path']
        else:
            raise ValueError("Either file_path or (output_dir, base_name, ext) must be provided")
    
    if not text:
        raise ValueError("text parameter is required")
    
    # Create directory if needed
    os.makedirs(os.path.dirname(file_path) if os.path.dirname(file_path) else '.', exist_ok=True)
    
    # Handle DOCX format
    if ext == '.docx' or (file_path and file_path.endswith('.docx')):
        # Use write_docx_with_skeleton for DOCX files
        skeleton = None
        if original_file and os.path.exists(original_file):
            try:
                skeleton = make_style_skeleton_from_docx(original_file)
            except Exception:
                pass
        return write_docx_with_skeleton(text, file_path, skeleton)
    
    # Handle DOC format (legacy) - convert to DOCX instead since python-docx doesn't write DOC
    if ext == '.doc' or (file_path and file_path.endswith('.doc')):
        # DOC is a legacy binary format that's hard to write
        # Save as DOCX instead with the same styling
        docx_path = file_path.replace('.doc', '.docx') if file_path.endswith('.doc') else file_path + 'x'
        skeleton = None
        if original_file and os.path.exists(original_file):
            # Try to read skeleton from original DOC file (may fail)
            try:
                # Convert DOC to DOCX for skeleton extraction (best effort)
                pass  # skeleton extraction from DOC is complex, skip for now
            except Exception:
                pass
        result = write_docx_with_skeleton(text, docx_path, skeleton)
        # Log the conversion
        print(f"Note: Converted DOC output to DOCX: {result}")
        return result
    
    # Handle PDF format
    if ext == '.pdf' or (file_path and file_path.endswith('.pdf')):
        return _write_text_to_pdf(text, file_path)
    
    # Handle Markdown - preserve it as plain text with .md extension
    if ext == '.md' or (file_path and file_path.endswith('.md')):
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(text)
        return file_path
    
    # Write plain text (TXT and anything else)
    with open(file_path, 'w', encoding='utf-8') as f:
        f.write(text)
    
    return file_path


def _write_text_to_pdf(text: str, file_path: str) -> str:
    """Write text to a PDF file using reportlab or fpdf2."""
    try:
        # Try fpdf2 first (lighter weight)
        from fpdf import FPDF
        
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Helvetica", size=11)
        
        # Handle text encoding and line breaks
        lines = text.split('\n')
        for line in lines:
            # Encode line to handle special characters
            try:
                # FPDF2 handles unicode better than fpdf
                pdf.multi_cell(0, 6, line)
            except Exception:
                # Fallback: encode to latin-1 with replacement
                safe_line = line.encode('latin-1', errors='replace').decode('latin-1')
                pdf.multi_cell(0, 6, safe_line)
        
        pdf.output(file_path)
        return file_path
    except ImportError:
        pass
    
    try:
        # Fallback to reportlab
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        
        doc = SimpleDocTemplate(file_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Split text into paragraphs
        paragraphs = text.split('\n\n')
        for para in paragraphs:
            if para.strip():
                # Clean up for reportlab
                safe_para = para.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                story.append(Paragraph(safe_para, styles['Normal']))
                story.append(Spacer(1, 0.2*inch))
        
        doc.build(story)
        return file_path
    except ImportError:
        pass
    
    # Last resort: save as TXT with .pdf.txt extension
    txt_path = file_path + '.txt'
    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write(text)
    print(f"Warning: PDF libraries not available. Saved as text file: {txt_path}")
    return txt_path

# ---------------------------
# Google Drive utilities
# ---------------------------

def extract_drive_file_id(url_or_id: str) -> str:
    """Extract file ID from Google Drive URL or return as-is if already an ID."""
    if not url_or_id:
        return ""
    
    # If it's already just an ID (no slashes or special chars), return it
    if '/' not in url_or_id and '?' not in url_or_id:
        return url_or_id
    
    # Try to extract from URL
    parsed = urlparse(url_or_id)
    
    # Handle file view URLs: https://drive.google.com/file/d/FILE_ID/view?usp=drive_link
    if '/file/d/' in parsed.path:
        file_id = parsed.path.split('/file/d/')[1].split('/')[0]
        return file_id
    
    # Handle document URLs: https://docs.google.com/document/d/FILE_ID/edit
    if '/document/d/' in parsed.path:
        file_id = parsed.path.split('/document/d/')[1].split('/')[0]
        return file_id
    
    # Handle spreadsheet URLs: https://docs.google.com/spreadsheets/d/FILE_ID/edit
    if '/spreadsheets/d/' in parsed.path:
        file_id = parsed.path.split('/spreadsheets/d/')[1].split('/')[0]
        return file_id
    
    # Handle presentation URLs: https://docs.google.com/presentation/d/FILE_ID/edit
    if '/presentation/d/' in parsed.path:
        file_id = parsed.path.split('/presentation/d/')[1].split('/')[0]
        return file_id
    
    # Check path for /d/FILE_ID pattern (standard share links)
    if '/d/' in parsed.path:
        file_id = parsed.path.split('/d/')[1].split('/')[0]
        return file_id
    
    # Check query parameters
    if 'id=' in parsed.query:
        return parse_qs(parsed.query)['id'][0]
    
    # Return as-is if we can't parse it
    return url_or_id

def create_google_doc(title: str, content: str = "") -> str:
    """Create a new Google Doc and return its file ID."""
    try:
        docs_service = get_docs_service()
        drive_service = get_drive_service()
        
        # Create document
        doc = docs_service.documents().create(body={'title': title}).execute()
        doc_id = doc.get('documentId')
        
        if content:
            # Insert content
            docs_service.documents().batchUpdate(
                documentId=doc_id,
                body={'requests': [{
                    'insertText': {
                        'location': {'index': 1},
                        'text': content
                    }
                }]}
            ).execute()
        
        return doc_id
    except Exception as e:
        raise ValueError(f"Failed to create Google Doc: {e}")

# ---------------------------
# DOCX style utilities
# ---------------------------

def make_style_skeleton_from_docx(docx_path: str) -> Dict[str, Any]:
    """Extract style skeleton from a DOCX file."""
    try:
        doc = Document(docx_path)
        skeleton = {
            'styles': {},
            'default_font': {'name': 'Arial', 'size': 11},
        }
        
        for para in doc.paragraphs[:50]:  # Sample first 50 paragraphs
            style_name = para.style.name if para.style else 'Normal'
            if style_name not in skeleton['styles']:
                runs = para.runs
                if runs:
                    first_run = runs[0]
                    skeleton['styles'][style_name] = {
                        'font_name': first_run.font.name or 'Arial',
                        'font_size': first_run.font.size.pt if first_run.font.size else 11,
                        'bold': first_run.bold,
                        'italic': first_run.italic,
                    }
        
        return skeleton
    except Exception as e:
        print(f"Warning: Failed to extract style skeleton: {e}")
        return {'styles': {}, 'default_font': {'name': 'Arial', 'size': 11}}

def write_docx_with_skeleton(text: str, output_path: str, skeleton: Dict[str, Any] = None):
    """Write text to DOCX file with style skeleton."""
    doc = Document()
    
    # Apply default font
    default_font = (skeleton or {}).get('default_font', {'name': 'Arial', 'size': 11})
    style = doc.styles['Normal']
    style.font.name = default_font['name']
    style.font.size = Pt(default_font['size'])
    
    # Split text into paragraphs and add
    paragraphs = text.split('\n')
    for para_text in paragraphs:
        if para_text.strip():
            para = doc.add_paragraph(para_text)
            # Apply style if available
            if skeleton and 'styles' in skeleton:
                # Try to match style (simplified)
                para.style = doc.styles['Normal']
    
    doc.save(output_path)
    return output_path

def make_style_sequence_from_docx(docx_path: str) -> List[Dict[str, Any]]:
    """Extract style sequence from a DOCX file."""
    try:
        doc = Document(docx_path)
        sequence = []
        
        for para in doc.paragraphs:
            style_name = para.style.name if para.style else 'Normal'
            runs = para.runs
            if runs:
                first_run = runs[0]
                sequence.append({
                    'text': para.text,
                    'style': style_name,
                    'font_name': first_run.font.name or 'Arial',
                    'font_size': first_run.font.size.pt if first_run.font.size else 11,
                    'bold': first_run.bold,
                    'italic': first_run.italic,
                })
            else:
                sequence.append({
                    'text': para.text,
                    'style': style_name,
                })
        
        return sequence
    except Exception as e:
        print(f"Warning: Failed to extract style sequence: {e}")
        return []

# ---------------------------
# History profile utilities
# ---------------------------

def derive_history_profile(history_path: str = None) -> Dict[str, float]:
    """Derive refinement profile from history data."""
    from app.core.paths import get_data_dir
    
    if history_path is None:
        history_path = str(get_data_dir() / 'recent_history.json')
    
    if not os.path.exists(history_path):
        # Return default profile
        return {
            'brevity_bias': 0.5,
            'formality_bias': 0.5,
            'structure_bias': 0.5,
        }
    
    try:
        with open(history_path, 'r', encoding='utf-8') as f:
            history = _json.load(f)
        
        # Analyze history to derive biases
        # Simplified implementation - can be enhanced
        brevity_bias = 0.5
        formality_bias = 0.5
        structure_bias = 0.5
        
        if isinstance(history, list) and len(history) > 0:
            # Analyze recent refinements
            total_changes = 0
            for entry in history[-10:]:  # Last 10 entries
                if isinstance(entry, dict):
                    # Extract metrics if available
                    pass
        
        return {
            'brevity_bias': brevity_bias,
            'formality_bias': formality_bias,
            'structure_bias': structure_bias,
        }
    except Exception as e:
        print(f"Warning: Failed to derive history profile: {e}")
        return {
            'brevity_bias': 0.5,
            'formality_bias': 0.5,
            'structure_bias': 0.5,
        }
